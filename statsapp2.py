import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
from PyPDF2 import PdfReader
import re

# ---------------- Streamlit page config ----------------

st.set_page_config(page_title="College Basketball Stats Organizer", layout="wide")
st.title("KenPom & CBBAnalytics Stats Organizer")
st.write(
    "Copy the **advanced stats table directly from KenPom**, paste the raw text below, "
    "and the app will parse it into clean advanced stats.\n\n"
    "Upload a **CBB Player Stats CSV** "
    "to generate a OVERVIEW DOCX.\n\n"
    "Upload a **CBB Player Profiles PDF** "
    "to generate a Shot Diet and FG% DOCX.\n\n"
    "The app sorts each category, adds ranks, colors titles and headers with your logo color, "
    "and exports Word docs in a two-column format.\n\n"
    "**Team name and team logo are required before generating any DOCX.**"
)

# ---------------- Helpers ----------------

def clean_player_name(name: str) -> str:
    """
    Cleans KenPom names by removing junk like 'National Rank' etc.
    """
    if not isinstance(name, str):
        name = str(name)

    s = name.replace("\n", " ")
    # Remove 'National Rank...' (case-insensitive, optional spaces)
    s = re.sub(r"(?i)national\s*rank.*$", "", s)
    # Remove trailing digits like '1' in 'Boozer1'
    s = re.sub(r"\d+$", "", s)
    # Collapse whitespace
    s = " ".join(s.split())
    return s.strip()


def clean_jersey(j):
    """
    Convert jersey values like '1', '01', '1.0', '1.00', ' 1.0' → '1'.
    If it's not numeric, return stripped string.
    """
    try:
        return str(int(float(str(j).strip())))
    except Exception:
        return str(j).strip()


def _final_clean_kenpom_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Shared final cleaning logic for KenPom advanced stats DataFrame.
    Assumes df already has 'Jersey' and 'Player' columns populated.
    """
    # Remove category header rows (blank jerseys)
    df = df[df["Jersey"].notna() & (df["Jersey"].str.len() > 0)].copy()

    # Clean jersey formatting
    df["Jersey"] = (
        df["Jersey"]
        .astype(str)
        .str.replace(".0", "", regex=False)
        .str.strip()
    )

    # Clean player names (remove National Rank junk)
    df["Player"] = df["Player"].apply(clean_player_name)

    allowed_columns = [
        "Jersey",
        "Player",
        "ORtg",
        "%Poss",
        "%Shots",
        "eFG%",
        "TS%",
        "OR%",
        "DR%",
        "ARate",
        "TORate",
        "Blk%",
        "Stl%",
        "FC/40",
        "FD/40",
        "FTRate",
    ]

    df = df[[c for c in df.columns if c in allowed_columns]].copy()

    numeric_cols = [c for c in df.columns if c not in ["Jersey", "Player"]]

    for col in numeric_cols:
        df[col] = (
            df[col]
            .astype(str)
            .str.replace("%", "", regex=False)
            .str.replace(",", "", regex=False)
            .str.strip()
        )
        df[col] = pd.to_numeric(df[col], errors="coerce")

        # Truncate to 1 decimal place (KenPom style)
        df[col] = df[col].apply(
            lambda x: float(int(x * 10)) / 10 if pd.notna(x) else x
        )

    return df


def load_and_clean_kenpom_csv(uploaded_file):
    """
    Fallback: Loads a KenPom-style advanced stats CSV:
      - First two columns are Unnamed: 0 (jersey), Unnamed: 1 (name), etc.
    """
    df = pd.read_csv(uploaded_file)

    # Clean whitespace from headers
    df.columns = [c.strip() for c in df.columns]

    # Jersey + player
    if "Unnamed: 0" in df.columns and "Unnamed: 1" in df.columns:
        df["Jersey"] = df["Unnamed: 0"].astype(str).str.strip()
        df["Player"] = df["Unnamed: 1"].astype(str).str.strip()
        df = df.drop(columns=["Unnamed: 0", "Unnamed: 1"])
    else:
        # Fallback if headers were renamed
        jersey_col = None
        player_col = None
        for c in df.columns:
            lc = c.lower()
            if "jersey" in lc or lc in ("#", "no", "number"):
                jersey_col = c
            if "player" in lc or "name" in lc:
                player_col = c

        if jersey_col is None or player_col is None:
            raise ValueError("Could not find jersey/name columns in advanced stats CSV.")

        df["Jersey"] = df[jersey_col].astype(str).str.strip()
        df["Player"] = df[player_col].astype(str).str.strip()

    df = _final_clean_kenpom_df(df)
    return df


def extract_numbers(line: str):
    """Return list of floatable numbers found in a line."""
    nums = re.findall(r"[-+]?\d*\.?\d+", line)
    out = []
    for n in nums:
        try:
            out.append(float(n))
        except ValueError:
            continue
    return out


def parse_kenpom_paste(raw: str) -> pd.DataFrame:
    """
    Parse raw text copied directly from a KenPom player-usage/advanced page.
    """

    import re

    # ---- Clean lines ----
    lines = [ln.rstrip() for ln in raw.splitlines()]
    lines = [ln for ln in lines if ln.strip()]

    # Player blocks start on lines like: "20 Michael McNair"
    player_starts = []
    for idx, line in enumerate(lines):
        if re.match(r"^\s*\d+\s+[A-Za-z]", line):
            player_starts.append(idx)

    if not player_starts:
        raise ValueError("No player lines found in pasted text.")

    players = []

    def parse_player_block(block_lines):
        text = " ".join(block_lines)
        tokens = text.split()
        if not tokens or not tokens[0].isdigit():
            return None

        jersey = tokens[0]

        # --- Name up to height/year ---
        name_tokens = []
        height_idx = None
        year_tokens = {"Fr", "So", "Jr", "Sr", "Gr", "Fr.", "So.", "Jr.", "Sr."}

        for i in range(1, len(tokens)):
            t = tokens[i]
            if re.match(r"^\d+-\d+$", t) or t in year_tokens:
                if re.match(r"^\d+-\d+$", t):
                    height_idx = i
                break
            else:
                name_tokens.append(t)

        if not name_tokens:
            return None

        # Clean name for pasted KenPom
        name = " ".join(name_tokens)
        name = re.sub(r"(?i)\s*national\s*rank.*$", "", name)  # drop 'National Rank'
        name = re.sub(r"\s*\d+$", "", name)                    # drop trailing digits like '1'
        name = " ".join(name.split())

        # If we didn't see height yet, scan for it
        if height_idx is None:
            for i in range(len(name_tokens) + 1, len(tokens)):
                t = tokens[i]
                if re.match(r"^\d+-\d+$", t):
                    height_idx = i
                    break

        if height_idx is None or height_idx + 2 >= len(tokens):
            # Still return jersey + player so at least something shows
            return {"Jersey": jersey, "Player": name}

        # Ht, Wt, Yr
        yr_idx = height_idx + 2

        # --- Detect %Min and ORtg robustly ---
        ortg = None
        ortg_token_index = None

        # Scan for first decimal 0–100 after Yr (= %Min)
        # Then scan forward for the next decimal 20–200 (= ORtg),
        # allowing integer ranks in between.
        for i in range(yr_idx + 1, len(tokens)):
            t = tokens[i]
            if "." not in t:
                continue
            try:
                v = float(t)
            except ValueError:
                continue
            if not (0.0 <= v <= 100.0):
                continue  # not %Min

            # Candidate %Min found -> look ahead for ORtg
            for j in range(i + 1, min(i + 8, len(tokens))):
                tj = tokens[j]
                if "." not in tj:
                    continue
                try:
                    v2 = float(tj)
                except ValueError:
                    continue
                if 20.0 <= v2 <= 200.0:  # ORtg range
                    ortg = v2
                    ortg_token_index = j
                    break
            if ortg_token_index is not None:
                break

        if ortg_token_index is None:
            # Couldn't confidently find ORtg; skip this player
            return None

        # --- Advanced zone: everything after ORtg up to splits (8-9, 11-13, etc.) ---
        adv_start = ortg_token_index + 1
        adv_end = len(tokens)
        for j in range(adv_start, len(tokens)):
            t = tokens[j]
            if "-" in t or t in ("FTM-A", "2PM-A", "3PM-A"):
                adv_end = j
                break

        adv_tokens = tokens[adv_start:adv_end]

        # KenPom rule: real stats have a decimal; ranks are plain ints.
        stat_tokens = [t for t in adv_tokens if "." in t]

        # First 13 decimals → advanced stats in order:
        # %Poss, %Shots, eFG%, TS%, OR%, DR%, ARate,
        # TORate, Blk%, Stl%, FC/40, FD/40, FTRate
        stat_tokens = stat_tokens[:13]
        stat_values = []
        for t in stat_tokens:
            try:
                stat_values.append(float(t))
            except ValueError:
                stat_values.append(None)

        while len(stat_values) < 13:
            stat_values.append(None)

        stat_cols_order = [
            "ORtg",
            "%Poss",
            "%Shots",
            "eFG%",
            "TS%",
            "OR%",
            "DR%",
            "ARate",
            "TORate",
            "Blk%",
            "Stl%",
            "FC/40",
            "FD/40",
            "FTRate",
        ]

        row = {"Jersey": jersey, "Player": name, "ORtg": ortg}
        for col_name, val in zip(stat_cols_order[1:], stat_values):
            row[col_name] = val

        return row

    # Build blocks and parse
    for idx, start_idx in enumerate(player_starts):
        end_idx = player_starts[idx + 1] if idx + 1 < len(player_starts) else len(lines)
        row = parse_player_block(lines[start_idx:end_idx])
        if row is not None:
            players.append(row)

    if not players:
        raise ValueError("Could not parse any players from the pasted KenPom text.")

    df = pd.DataFrame(players)
    df = _final_clean_kenpom_df(df)
    return df


def load_and_clean_overview_csv(uploaded_file):
    """
    Loads a CBB Analytics-style overview CSV and returns:
      Jersey, Player, and the overview stats used in the layout.
    Includes 'fga' (total FGA) for FG% makes/attempts logic.
    """
    df = pd.read_csv(uploaded_file)

    # Standardize to our naming
    df = df.rename(
        columns={
            "fullName": "Player",
            "jerseyNum": "Jersey",
        }
    )

    df["Jersey"] = df["Jersey"].apply(clean_jersey)
    df["Player"] = df["Player"].astype(str).str.strip()

    needed_cols = [
        "Jersey",
        "Player",
        "tsPct",
        "fgaP40",
        "fg2Pct",
        "astPct",
        "astTov",
        "tovPct",
        "astUsage",
        "drbPct",
        "blkPct",
        "fg3Pct",
        "ftPct",
        "fga3Rate",
        "usagePct",
        "ftaRate",
        "orbPct",
        "stlPct",
        "pfP40",
        "pfEff",
        "fga",  # total FGA for season
    ]

    cols_present = [c for c in needed_cols if c in df.columns]
    df = df[cols_present].copy()

    for col in df.columns:
        if col in ["Jersey", "Player"]:
            continue
        df[col] = pd.to_numeric(df[col], errors="coerce")

    return df


def remove_table_borders(table):
    """
    Remove all visible borders from a python-docx table.
    """
    tbl = table._element

    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    border_names = ["top", "left", "bottom", "right", "insideH", "insideV"]

    for border_name in border_names:
        border = tblBorders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tblBorders.append(border)
        border.set(qn("w:val"), "nil")


def format_overview_value(col: str, value: float) -> str:
    """
    Convert overview stats into correct display formats.
    """
    if pd.isna(value):
        return ""

    # Ratio stats
    if col in ["astTov", "astUsage", "pfEff"]:
        return f"{value:.2f}x"

    # Per-40 stats
    if col in ["fgaP40", "pfP40"]:
        return f"{value:.1f}"

    # Percent-like stats
    percent_cols = [
        "tsPct", "fg2Pct", "fg3Pct", "ftPct",
        "fga3Rate", "usagePct", "ftaRate","blkPct"
        "orbPct", "drbPct", "stlPct", "tovPct","astPct"
    ]
    if col in percent_cols:
        if value < 10:
            value = value * 100
        return f"{value:.1f}%"

    return f"{value:.1f}"


def get_darkest_color_from_logo(logo_bytes):
    """
    Given raw logo bytes, return an RGBColor corresponding to a dark
    NON-BACKGROUND color in the image.
    """
    try:
        img = Image.open(BytesIO(logo_bytes)).convert("RGBA")
    except Exception:
        return None

    img = img.resize((64, 64))

    def find_dark_color(ignore_near_black: bool):
        darkest_rgb = None
        darkest_lum = float("inf")

        for r, g, b, a in img.getdata():
            if a == 0:
                continue

            if ignore_near_black and r < 15 and g < 15 and b < 15:
                continue

            lum = 0.2126 * r + 0.7152 * g + 0.0722 * b
            if lum < darkest_lum:
                darkest_lum = lum
                darkest_rgb = (r, g, b)

        return darkest_rgb

    darkest_rgb = find_dark_color(ignore_near_black=True)
    if darkest_rgb is None:
        darkest_rgb = find_dark_color(ignore_near_black=False)

    if darkest_rgb is None:
        return None

    return RGBColor(darkest_rgb[0], darkest_rgb[1], darkest_rgb[2])


def parse_cbb_team_pdf_shooting_zones(uploaded_pdf) -> pd.DataFrame:
    """
    Given a CBB Analytics TEAM player-profiles PDF, extract FGA% by shot zone
    from each player's 'Shot Zone GP* FGA/G FGA% FG%' table on their page.

    Returns a DataFrame with:
      Jersey, Player,
      FGA% At Rim, FGA% In Paint, FGA% Midrange 2s,
      FGA% Above Break 3s, FGA% Corner 3s
    """
    from io import BytesIO
    from PyPDF2 import PdfReader

    pdf_bytes = uploaded_pdf.read()
    reader = PdfReader(BytesIO(pdf_bytes))

    rows = []
    zones = [
        "At Rim",
        "In Paint",
        "Midrange 2s",
        "Above Break 3s",
        "Corner 3s",
        "At Rim + 3s",
        "Heaves",
    ]

    for page in reader.pages:
        try:
            text = page.extract_text()
        except Exception:
            continue

        if not text or "Shot Zone" not in text or "FGA%" not in text:
            continue

        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

        # Header line looks like: "DJ Richards (#2, Guard, 6'4") : Profile"
        header_line = lines[0]
        name = header_line
        jersey = ""

        m = re.match(r"^(.*?)\s*\(#(\d+)", header_line)
        if m:
            name = m.group(1).strip()
            jersey = m.group(2).strip()
        else:
            # fallback: name before "(" if no jersey matched
            name = header_line.split("(")[0].strip()

        # Find the "Shot Zone ... FGA%" line
        shot_idx = None
        for i, ln in enumerate(lines):
            if "Shot Zone" in ln and "FGA%" in ln:
                shot_idx = i
                break

        if shot_idx is None:
            continue

        # Now grab the rows that follow the Shot Zone header
        fga_vals = []
        for ln in lines[shot_idx + 1:]:
            # Stop if we hit another section
            if ln.startswith("DNQ") or "Zone % of Shots" in ln or "Shot Chart" in ln:
                break

            parts = ln.split()
            if len(parts) < 2:
                continue

            # e.g., "1.0 33.3% 50.0%"
            fga_token = parts[1]
            if "%" not in fga_token:
                continue

            try:
                v = float(fga_token.replace("%", ""))
            except ValueError:
                v = float("nan")

            fga_vals.append(v)
            if len(fga_vals) >= len(zones):
                break

        if not fga_vals:
            continue

        # Ensure we have exactly 7 values (pad if short)
        while len(fga_vals) < len(zones):
            fga_vals.append(float("nan"))

        row = {
            "Jersey": jersey,
            "Player": name,
        }

        # Assign each zone's FGA%
        for i, z in enumerate(zones):
            row[f"FGA% {z}"] = fga_vals[i]

        rows.append(row)

    if not rows:
        raise ValueError("No 'Shot Zone' tables found in the uploaded PDF.")

    df = pd.DataFrame(rows)

    # Only keep the 5 zones you care about
    keep_cols = [
        "Jersey",
        "Player",
        "FGA% At Rim",
        "FGA% In Paint",
        "FGA% Midrange 2s",
        "FGA% Above Break 3s",
        "FGA% Corner 3s",
    ]

    return df[keep_cols]


# ---------------- Streamlit UI ----------------

team_name = st.text_input(
    "Team name for DOCX titles (REQUIRED)",
    value="",
    help="Shown as '<TEAM> ADVANCED STATISTICS' and '<TEAM> OVERVIEW STATISTICS' in the documents.",
)

if team_name.strip() == "":
    st.error("❗ Team name is required.")

logo_file = st.file_uploader(
    "Upload Team Logo (REQUIRED)",
    type=["png", "jpg", "jpeg"],
    key="logo",
    help="Logo will appear at the top of both DOCX files, and its darkest non-black color will be used for titles and headers.",
)

if logo_file is None:
    st.error("❗ Team logo is required.")

title_text = (team_name or "").strip().upper()
safe_team_name = (title_text or "TEAM").replace(" ", "_")

logo_bytes = None
header_color = None
if logo_file is not None:
    logo_bytes = logo_file.read()
    if logo_bytes:
        header_color = get_darkest_color_from_logo(logo_bytes)

# =========================================================
#  ADVANCED STATS INPUT SECTION (KENPOM)
# =========================================================

st.markdown("### Advanced Stats Input (KenPom)")

pasted_kenpom = st.text_area(
    "Paste raw KenPom advanced/usage table here:",
    height=300,
    help="Highlight the whole advanced/usage table on KenPom → Copy → Paste here.",
)

uploaded_file = st.file_uploader(
    "Or upload an Advanced Stats CSV",
    type=["csv"],
    help="CSV must include: ORtg, %Poss, %Shots, eFG%, TS%, OR%, DR%, ARate, TORate, Blk%, Stl%, FC/40, FD/40, FTRate",
)

# ---------------- BLUE NOTE FOR ADVANCED STATS ----------------
st.markdown(
    """
<div style="
    background-color:#0A66C2;
    padding:12px;
    border-radius:6px;
    color:white;
    font-size:15px;
    margin-top:10px;
">
⬆️ Paste KenPom text above or upload an <strong>Advanced Stats CSV</strong> to generate the Advanced Stats DOCX
(after entering team name and logo).
</div>
""",
    unsafe_allow_html=True
)

# ---------- ADVANCED STATS PIPELINE ----------

df_stats = None

if pasted_kenpom.strip():
    try:
        df_stats = parse_kenpom_paste(pasted_kenpom)
    except Exception as e:
        st.error(f"Error parsing pasted KenPom text: {e}")
        df_stats = None
elif uploaded_file is not None:
    try:
        df_stats = load_and_clean_kenpom_csv(uploaded_file)
    except Exception as e:
        st.error(f"Error reading advanced stats CSV: {e}")
        df_stats = None

if df_stats is not None:
    if title_text == "":
        st.error("❗ Please enter a team name before generating the Advanced Stats DOCX.")
    elif logo_bytes is None:
        st.error("❗ Please upload a team logo before generating the Advanced Stats DOCX.")
    else:
        st.subheader("Parsed Advanced Stats Table")
        st.dataframe(df_stats, use_container_width=True)

        categories = [
            ("ORtg", "Offensive Rating"),
            ("FTRate", "Free-Throw Rate"),
            ("%Poss", "% of Possessions"),
            ("eFG%", "Effective FG%"),
            ("%Shots", "% of Shots"),
            ("TS%", "True Shooting %"),
            ("OR%", "Offensive Rebound %"),
            ("DR%", "Defensive Rebound %"),
            ("TORate", "Turnover Rate"),
            ("ARate", "Assist Rate"),
            ("FD/40", "Fouls Drawn per 40"),
            ("FC/40", "Fouls Committed per 40"),
            ("Blk%", "Block %"),
            ("Stl%", "Steal %"),
        ]

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        if logo_bytes:
            logo_stream = BytesIO(logo_bytes)
            doc.add_picture(logo_stream, width=Inches(1.2))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_paragraph = doc.add_paragraph()
        run = title_paragraph.add_run(f"{title_text} ADVANCED STATISTICS")
        run.bold = True
        run.font.size = Pt(14)
        if header_color is not None:
            run.font.color.rgb = header_color
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(4)

        pairs = []
        for i in range(0, len(categories), 2):
            if i + 1 < len(categories):
                pairs.append((categories[i], categories[i + 1]))
            else:
                pairs.append((categories[i], None))

        for left_cat, right_cat in pairs:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            remove_table_borders(table)

            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # LEFT CATEGORY
            if left_cat is not None:
                col, title = left_cat
                df_sorted = (
                    df_stats.sort_values(by=col, ascending=False)
                    if col in df_stats.columns
                    else None
                )

                p = left_cell.add_paragraph()
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(16)
                if header_color is not None:
                    r.font.color.rgb = header_color
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(11)

                if df_sorted is not None:
                    for rank, (_, row) in enumerate(df_sorted.iterrows(), start=1):
                        value = row[col]
                        if pd.isna(value):
                            continue
                        jersey = clean_jersey(row["Jersey"])
                        name = str(row["Player"])

                        if abs(value - int(value)) < 1e-6:
                            val_str = f"{int(value)}"
                        else:
                            val_str = f"{value:.1f}"

                        suffix = "" if col in ["ORtg", "FC/40", "FD/40"] else "%"

                        pl = left_cell.add_paragraph()
                        pr = pl.add_run(f"{rank}. #{jersey} {name} – {val_str}{suffix}")
                        pr.font.size = Pt(11)
                        pl.paragraph_format.space_before = Pt(0)
                        pl.paragraph_format.space_after = Pt(0)
                        pl.paragraph_format.line_spacing = Pt(11)

            # RIGHT CATEGORY
            if right_cat is not None:
                col, title = right_cat
                df_sorted = (
                    df_stats.sort_values(by=col, ascending=False)
                    if col in df_stats.columns
                    else None
                )

                p = right_cell.add_paragraph()
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(16)
                if header_color is not None:
                    r.font.color.rgb = header_color
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(11)

                if df_sorted is not None:
                    for rank, (_, row) in enumerate(df_sorted.iterrows(), start=1):
                        value = row[col]
                        if pd.isna(value):
                            continue
                        jersey = clean_jersey(row["Jersey"])
                        name = str(row["Player"])

                        if abs(value - int(value)) < 1e-6:
                            val_str = f"{int(value)}"
                        else:
                            val_str = f"{value:.1f}"

                        suffix = "" if col in ["ORtg", "FC/40", "FD/40"] else "%"

                        pl = right_cell.add_paragraph()
                        pr = pl.add_run(f"{rank}. #{jersey} {name} – {val_str}{suffix}")
                        pr.font.size = Pt(11)
                        pl.paragraph_format.space_before = Pt(0)
                        pl.paragraph_format.space_after = Pt(0)
                        pl.paragraph_format.line_spacing = Pt(11)

            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.line_spacing = Pt(0.25)

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        filename = f"{safe_team_name}_advanced_statistics.docx"

        st.download_button(
            label="Download Advanced Stats DOCX",
            data=docx_buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_advanced",
        )

    



st.markdown("---")  # Divider between sections


# =========================================================
#  OVERVIEW STATS INPUT SECTION (CBB)
# =========================================================

st.markdown("### Overview Stats Input (CBB)")

overview_file = st.file_uploader(
    "Upload a CBB Overview CSV (TS%, 3P%, FGA/40, FT%, 2P%, 3PAr, AST%, USG%, AST/USG, AST/TOV, FTAr, TOV%, ORB%, DRB%, STL%, PF/40, BLK%, PF Eff, FGA)",
    type=["csv"],
    help="This will be used to generate the Overview DOCX and to compute FG% makes/attempts.",
)

# ---------------- BLUE NOTE FOR OVERVIEW STATS ----------------
st.markdown(
    """
<div style="
    background-color:#0A66C2;
    padding:12px;
    border-radius:6px;
    color:white;
    font-size:15px;
    margin-top:10px;
">
⬆️ Upload a <strong>CBB Overview CSV</strong> to generate the Overview DOCX
(after entering team name and logo).
</div>
""",
    unsafe_allow_html=True
)

st.markdown("---")

# ---------- OVERVIEW PIPELINE ----------

if overview_file is not None:
    if title_text == "":
        st.error("❗ Please enter a team name before generating the Overview DOCX.")
    elif logo_bytes is None:
        st.error("❗ Please upload a team logo before generating the Overview DOCX.")
    else:
        try:
            overview_file.seek(0)
            df_overview = load_and_clean_overview_csv(overview_file)
        except Exception as e:
            st.error(f"Error reading overview CSV: {e}")
            st.stop()

        st.subheader("Parsed Overview Table")
        st.dataframe(df_overview, use_container_width=True)

        left_overview = [
            ("tsPct", "True Shooting %"),
            ("fgaP40", "Field Goal Attempts per 40 (FGA/40)"),
            ("fg2Pct", "Two-Point %"),
            ("astPct", "Assist %"),
            ("astTov", "Assist-to-Turnover"),
            ("tovPct", "Turnover %"),
            ("astUsage", "Assist/Usage"),
            ("drbPct", "Defensive Rebound %"),
            ("blkPct", "Block %"),
        ]

        right_overview = [
            ("fg3Pct", "Three-Point %"),
            ("ftPct", "Free Throw %"),
            ("fga3Rate", "Three-Point Attempt Rate"),
            ("usagePct", "Usage %"),
            ("ftaRate", "Free Throw Attempt Rate"),
            ("orbPct", "Offensive Rebound %"),
            ("stlPct", "Steal %"),
            ("pfP40", "Personal Fouls per 40"),
            ("pfEff", "PF Efficiency"),
        ]

        overview_pairs = []
        max_len = max(len(left_overview), len(right_overview))
        for i in range(max_len):
            left = left_overview[i] if i < len(left_overview) else None
            right = right_overview[i] if i < len(right_overview) else None
            overview_pairs.append((left, right))

        overview_doc = Document()
        style = overview_doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        if logo_bytes:
            logo_stream = BytesIO(logo_bytes)
            overview_doc.add_picture(logo_stream, width=Inches(1.2))
            last_paragraph = overview_doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_paragraph = overview_doc.add_paragraph()
        run = title_paragraph.add_run(f"{title_text} OVERVIEW STATISTICS")
        run.bold = True
        run.font.size = Pt(14)
        if header_color is not None:
            run.font.color.rgb = header_color
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(4)

        for left_cat, right_cat in overview_pairs:
            table = overview_doc.add_table(rows=1, cols=2)
            table.autofit = True
            remove_table_borders(table)

            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # LEFT
            if left_cat is not None:
                col, title = left_cat
                p = left_cell.add_paragraph()
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(16)
                if header_color is not None:
                    r.font.color.rgb = header_color
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(11)

                if col in df_overview.columns:
                    df_sorted = df_overview.sort_values(by=col, ascending=False)
                    for rank, (_, row) in enumerate(df_sorted.iterrows(), start=1):
                        value = row[col]
                        if pd.isna(value):
                            continue
                        jersey = clean_jersey(row["Jersey"])
                        name = str(row["Player"])
                        val_str = format_overview_value(col, value)

                        pl = left_cell.add_paragraph()
                        pr = pl.add_run(f"{rank}. #{jersey} {name} – {val_str}")
                        pr.font.size = Pt(11)
                        pl.paragraph_format.space_before = Pt(0)
                        pl.paragraph_format.space_after = Pt(0)
                        pl.paragraph_format.line_spacing = Pt(11)

            # RIGHT
            if right_cat is not None:
                col, title = right_cat
                p = right_cell.add_paragraph()
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(16)
                if header_color is not None:
                    r.font.color.rgb = header_color
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(11)

                if col in df_overview.columns:
                    df_sorted = df_overview.sort_values(by=col, ascending=False)
                    for rank, (_, row) in enumerate(df_sorted.iterrows(), start=1):
                        value = row[col]
                        if pd.isna(value):
                            continue
                        jersey = clean_jersey(row["Jersey"])
                        name = str(row["Player"])
                        val_str = format_overview_value(col, value)

                        pl = right_cell.add_paragraph()
                        pr = pl.add_run(f"{rank}. #{jersey} {name} – {val_str}")
                        pr.font.size = Pt(11)
                        pl.paragraph_format.space_before = Pt(0)
                        pl.paragraph_format.space_after = Pt(0)
                        pl.paragraph_format.line_spacing = Pt(11)

            sp = overview_doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.line_spacing = Pt(0.25)

        overview_buffer = BytesIO()
        overview_doc.save(overview_buffer)
        overview_buffer.seek(0)

        overview_filename = f"{safe_team_name}_overview_statistics.docx"

        st.download_button(
            label="Download Overview DOCX",
            data=overview_buffer,
            file_name=overview_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_overview",
        )

# =========================================================
#  CBB Analytics – Shooting by Region (Team PDF)
# =========================================================

st.subheader("CBB Analytics – Shooting by Region (Team PDF)")

team_pdf = st.file_uploader(
    "Upload CBB Analytics *Team Player-profiles* PDF (with 'Shooting by Region (Full Season)' tables)",
    type=["pdf"],
    key="cbb_team_pdf",
)

if team_pdf is not None:
    if not title_text:
        st.error("❗ Please enter a team name before generating the Shot Diet and FG% DOCX files.")
    elif logo_bytes is None:
        st.error("❗ Please upload a team logo before generating the Shot Diet and FG% DOCX files.")
    else:
        # ================================================================
        #  PARSE SHOT ZONE TABLE ONCE: GET BOTH FGA% AND FG% BY ZONE
        # ================================================================
        from io import BytesIO
        import re

        pdf_bytes = team_pdf.read()  # read once

        def parse_cbb_team_pdf_zones_both(pdf_bytes: bytes) -> pd.DataFrame:
            """
            Parse each player page's 'Shot Zone GP* FGA/... FGA% FG%' table.

            For each player we return:
              Jersey, Player,
              FGA% At Rim / In Paint / Midrange 2s / Above Break 3s / Corner 3s / At Rim + 3s / Heaves
              FG%  At Rim / In Paint / Midrange 2s / Above Break 3s / Corner 3s / At Rim + 3s / Heaves
            """
            reader_local = PdfReader(BytesIO(pdf_bytes))
            rows = []
            zones = [
                "At Rim",
                "In Paint",
                "Midrange 2s",
                "Above Break 3s",
                "Corner 3s",
                "At Rim + 3s",
                "Heaves",
            ]

            for page in reader_local.pages:
                try:
                    text = page.extract_text()
                except Exception:
                    continue
                if not text or "Shot Zone" not in text:
                    continue

                lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
                if not lines:
                    continue

                # Player header: "Name (#12, Guard, ... ) : Profile"
                header_line = lines[0]
                m = re.match(r"^(.*?)\s*\(#(\d+)", header_line)
                if m:
                    name = m.group(1).strip()
                    jersey = m.group(2).strip()
                else:
                    # Not a standard player page
                    continue

                # Find the shot-zone header line
                shot_idx = None
                for i, ln in enumerate(lines):
                    if "Shot Zone" in ln and "FGA" in ln:
                        shot_idx = i
                        break
                if shot_idx is None:
                    continue

                fga_vals = []
                fg_vals = []

                for ln in lines[shot_idx + 1 :]:
                    # Stop when we hit DNQ / summary sections
                    if (
                        ln.startswith("DNQ")
                        or "Zone % of Shots" in ln
                        or "Zone FG%" in ln
                        or "Shot Chart" in ln
                    ):
                        break

                    parts = ln.split()
                    if not parts:
                        continue

                    # Grab ALL numeric-like tokens (with or without %)
                    numeric_tokens = [
                        t for t in parts
                        if re.match(r"^-?\d+(\.\d+)?%?$", t)
                    ]

                    # Need at least 2 numeric tokens to get FGA% and FG%
                    if len(numeric_tokens) < 2:
                        continue

                    # Last two numeric tokens on the line are FGA% and FG%
                    fga_token = numeric_tokens[-2]
                    fg_token = numeric_tokens[-1]

                    try:
                        fga_vals.append(float(fga_token.replace("%", "")))
                    except Exception:
                        fga_vals.append(float("nan"))

                    try:
                        fg_vals.append(float(fg_token.replace("%", "")))
                    except Exception:
                        fg_vals.append(float("nan"))

                    if len(fga_vals) >= len(zones):
                        break

                if not fga_vals:
                    continue

                # Pad to 7 zones
                while len(fga_vals) < len(zones):
                    fga_vals.append(float("nan"))
                    fg_vals.append(float("nan"))
                while len(fg_vals) < len(zones):
                    fg_vals.append(float("nan"))

                row = {"Jersey": jersey, "Player": name}
                for i, z in enumerate(zones):
                    row[f"FGA% {z}"] = fga_vals[i]
                    row[f"FG% {z}"] = fg_vals[i]
                rows.append(row)

            if not rows:
                raise ValueError("No Shot Zone tables with both FGA% and FG% found in PDF.")
            return pd.DataFrame(rows)

        try:
            df_zones = parse_cbb_team_pdf_zones_both(pdf_bytes)
        except Exception as e:
            st.error(f"Could not extract shooting-by-region FGA% + FG% from PDF: {e}")
        else:
            # ------------------------------------------------------------
            # Split into separate FGA% and FG% DataFrames
            # ------------------------------------------------------------
            core_zones = ["At Rim", "In Paint", "Midrange 2s", "Above Break 3s", "Corner 3s"]
            fga_cols = [f"FGA% {z}" for z in core_zones]
            fg_cols = [f"FG% {z}" for z in core_zones]

            # SHOT DIET (FGA%)
            df_shooting = df_zones[["Jersey", "Player"] + fga_cols].copy()
            for c in fga_cols:
                df_shooting[c] = pd.to_numeric(df_shooting[c], errors="coerce")

            df_shooting["FGA% All 2 PT Attempts"] = (
                df_shooting["FGA% At Rim"]
                + df_shooting["FGA% In Paint"]
                + df_shooting["FGA% Midrange 2s"]
            )
            df_shooting["FGA% All 3 PT Attempts"] = (
                df_shooting["FGA% Above Break 3s"] + df_shooting["FGA% Corner 3s"]
            )

            # FG% BY ZONE
            df_fg_zones = df_zones[["Jersey", "Player"] + fg_cols].copy()
            for c in fg_cols:
                df_fg_zones[c] = pd.to_numeric(df_fg_zones[c], errors="coerce")

            # ============================================================
            #                      SHOT DIET PREVIEW
            # ============================================================
            st.markdown("### **Shot Diet Table Preview (Full Season)**")

            preview_cols = [
                "Jersey",
                "Player",
                "FGA% At Rim",
                "FGA% In Paint",
                "FGA% Midrange 2s",
                "FGA% All 2 PT Attempts",
                "FGA% Above Break 3s",
                "FGA% Corner 3s",
                "FGA% All 3 PT Attempts",
            ]
            st.dataframe(df_shooting[preview_cols], use_container_width=True)

            # ============================================================
            #                  FG% PREVIEW (FROM SAME TABLE)
            # ============================================================
            st.markdown("### **FG% Table Preview (From Shot Zone FG%)**")
            fg_preview_cols = [
                "Jersey",
                "Player",
                "FG% At Rim",
                "FG% In Paint",
                "FG% Midrange 2s",
                "FG% Above Break 3s",
                "FG% Corner 3s",
            ]
            st.dataframe(df_fg_zones[fg_preview_cols], use_container_width=True)

            # ============================================================
            #          REQUIRE OVERVIEW CSV (WITH FGA) FOR MAKES/ATTEMPTS
            # ============================================================
            if overview_file is None:
                st.error(
                    "❗ Upload a CBB Overview CSV (with an 'fga' column) to generate the FG% DOCX with makes/attempts."
                )
            else:
                overview_file.seek(0)
                try:
                    df_overview_all = load_and_clean_overview_csv(overview_file)
                except Exception as e:
                    st.error(f"Error reading overview CSV for FGA: {e}")
                    df_overview_all = None

                if df_overview_all is None or "fga" not in df_overview_all.columns:
                    st.error(
                        "❗ The Overview CSV must include an 'fga' column (total FGA per player)."
                    )
                else:
                    df_fga = df_overview_all[["Jersey", "Player", "fga"]].copy()

                    # Merge FG% by zone + FGA% by zone + total FGA
                    df_fg = df_fg_zones.merge(
                        df_shooting[["Jersey", "Player"] + fga_cols],
                        on=["Jersey", "Player"],
                        how="left",
                    )
                    df_fg = df_fg.merge(df_fga, on=["Jersey", "Player"], how="left")

                    # ========================================================
                    #          COMPUTE ZONE MAKES / ATTEMPTS USING FGA
                    # ========================================================
                    zone_defs = [
                        ("At Rim",         "FGA% At Rim",         "FG% At Rim"),
                        ("In Paint",       "FGA% In Paint",       "FG% In Paint"),
                        ("Midrange 2s",    "FGA% Midrange 2s",    "FG% Midrange 2s"),
                        ("Above Break 3s", "FGA% Above Break 3s", "FG% Above Break 3s"),
                        ("Corner 3s",      "FGA% Corner 3s",      "FG% Corner 3s"),
                    ]

                    for zone_name, _, _ in zone_defs:
                        df_fg[f"{zone_name} Attempts"] = 0
                        df_fg[f"{zone_name} Makes"] = 0
                        df_fg[f"{zone_name} FG"] = 0.0

                    for idx, row in df_fg.iterrows():
                        total_fga = row.get("fga", 0)
                        if pd.isna(total_fga) or total_fga <= 0:
                            continue

                        for zone_name, fga_pct_col, fg_pct_col in zone_defs:
                            fga_pct = row.get(fga_pct_col)
                            fg_pct = row.get(fg_pct_col)

                            if pd.isna(fga_pct) or pd.isna(fg_pct):
                                continue

                            zone_attempts = round(total_fga * (fga_pct / 100.0))
                            zone_makes = round(zone_attempts * (fg_pct / 100.0))

                            df_fg.at[idx, f"{zone_name} Attempts"] = int(zone_attempts)
                            df_fg.at[idx, f"{zone_name} Makes"] = int(zone_makes)
                            df_fg.at[idx, f"{zone_name} FG"] = (
                                0.0
                                if zone_attempts == 0
                                else 100.0 * zone_makes / zone_attempts
                            )

                    # ========================================================
                    #      TOTAL 2PT & 3PT MAKES / ATTEMPTS / FG% PER PLAYER
                    # ========================================================
                    df_fg["Total 2pt Attempts"] = (
                        df_fg["At Rim Attempts"]
                        + df_fg["In Paint Attempts"]
                        + df_fg["Midrange 2s Attempts"]
                    )
                    df_fg["Total 2pt Makes"] = (
                        df_fg["At Rim Makes"]
                        + df_fg["In Paint Makes"]
                        + df_fg["Midrange 2s Makes"]
                    )
                    df_fg["Total 2pt FG"] = 0.0
                    mask_2 = df_fg["Total 2pt Attempts"] > 0
                    df_fg.loc[mask_2, "Total 2pt FG"] = (
                        100.0
                        * df_fg.loc[mask_2, "Total 2pt Makes"]
                        / df_fg.loc[mask_2, "Total 2pt Attempts"]
                    )

                    df_fg["Total 3pt Attempts"] = (
                        df_fg["Above Break 3s Attempts"] + df_fg["Corner 3s Attempts"]
                    )
                    df_fg["Total 3pt Makes"] = (
                        df_fg["Above Break 3s Makes"] + df_fg["Corner 3s Makes"]
                    )
                    df_fg["Total 3pt FG"] = 0.0
                    mask_3 = df_fg["Total 3pt Attempts"] > 0
                    df_fg.loc[mask_3, "Total 3pt FG"] = (
                        100.0
                        * df_fg.loc[mask_3, "Total 3pt Makes"]
                        / df_fg.loc[mask_3, "Total 3pt Attempts"]
                    )

                    # ========================================================
                    #              BUILD SHOT DIET DOCX  (CONDENSED)
                    # ========================================================
                    shooting_doc = Document()
                    style = shooting_doc.styles["Normal"]
                    font = style.font
                    font.name = "Calibri"
                    font.size = Pt(11)
                    style.paragraph_format.space_before = Pt(0)
                    style.paragraph_format.space_after = Pt(0)

                    if logo_bytes:
                        logo_stream = BytesIO(logo_bytes)
                        shooting_doc.add_picture(logo_stream, width=Inches(1.2))
                        shooting_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    title_paragraph = shooting_doc.add_paragraph()
                    run = title_paragraph.add_run(f"{title_text} SHOT DIET")
                    run.bold = True
                    run.font.size = Pt(14)
                    if header_color is not None:
                        run.font.color.rgb = header_color
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_paragraph.paragraph_format.space_after = Pt(4)

                    left_zone_categories = [
                        ("FGA% At Rim", "FGA% -- At the Rim"),
                        ("FGA% In Paint", "FGA% -- In the Paint"),
                        ("FGA% Midrange 2s", "FGA% -- Mid-Range 2s"),
                        ("FGA% All 2 PT Attempts", "FGA% -- All 2 PT Attempts"),
                    ]
                    right_zone_categories = [
                        ("FGA% Above Break 3s", "FGA% -- Above the Break 3s"),
                        ("FGA% Corner 3s", "FGA% -- Corner 3s"),
                        ("FGA% All 3 PT Attempts", "FGA% -- All 3 PT Attempts"),
                    ]

                    max_rows = max(len(left_zone_categories), len(right_zone_categories))

                    # ONE table for all Shot Diet sections
                    shot_table = shooting_doc.add_table(rows=max_rows, cols=2)
                    shot_table.autofit = True
                    remove_table_borders(shot_table)

                    for i in range(max_rows):
                        left_cell = shot_table.rows[i].cells[0]
                        right_cell = shot_table.rows[i].cells[1]

                        # LEFT side
                        if i < len(left_zone_categories):
                            col, ttl = left_zone_categories[i]
                            p = left_cell.paragraphs[0]
                            p.text = ""
                            r = p.add_run(ttl)
                            r.bold = True
                            r.font.size = Pt(16)
                            if header_color is not None:
                                r.font.color.rgb = header_color

                            if col in df_shooting.columns:
                                df_sorted = df_shooting.sort_values(by=col, ascending=False)
                                for rank, (_, row_s) in enumerate(df_sorted.iterrows(), start=1):
                                    val_str = f"{row_s[col]:.1f}%"
                                    jersey = str(row_s["Jersey"]).strip()
                                    name = str(row_s["Player"])
                                    left_cell.add_paragraph(
                                        f"{rank}. #{jersey} {name} – {val_str}"
                                    )

                        # RIGHT side
                        if i < len(right_zone_categories):
                            col, ttl = right_zone_categories[i]
                            p = right_cell.paragraphs[0]
                            p.text = ""
                            r = p.add_run(ttl)
                            r.bold = True
                            r.font.size = Pt(16)
                            if header_color is not None:
                                r.font.color.rgb = header_color

                            if col in df_shooting.columns:
                                df_sorted = df_shooting.sort_values(by=col, ascending=False)
                                for rank, (_, row_s) in enumerate(df_sorted.iterrows(), start=1):
                                    val_str = f"{row_s[col]:.1f}%"
                                    jersey = str(row_s["Jersey"]).strip()
                                    name = str(row_s["Player"])
                                    right_cell.add_paragraph(
                                        f"{rank}. #{jersey} {name} – {val_str}"
                                    )

                    # Force zero spacing for every paragraph inside the table
                    for row in shot_table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(0)

                    shooting_buffer = BytesIO()
                    shooting_doc.save(shooting_buffer)
                    shooting_buffer.seek(0)

                    st.download_button(
                        "Download Shot Diet DOCX",
                        data=shooting_buffer,
                        file_name=f"{safe_team_name}_shot_diet.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_shot_diet",
                    )

                    # ========================================================
                    #                BUILD 2-COLUMN FG% DOCX  (CONDENSED)
                    # ========================================================
                    fg_doc = Document()
                    fg_style = fg_doc.styles["Normal"]
                    fg_style.font.name = "Calibri"
                    fg_style.font.size = Pt(11)
                    fg_style.paragraph_format.space_before = Pt(0)
                    fg_style.paragraph_format.space_after = Pt(0)

                    if logo_bytes:
                        fg_doc.add_picture(BytesIO(logo_bytes), width=Inches(1.2))
                        fg_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    title_p2 = fg_doc.add_paragraph()
                    r = title_p2.add_run(f"{title_text} FG%")
                    r.bold = True
                    r.font.size = Pt(14)
                    if header_color is not None:
                        r.font.color.rgb = header_color
                    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_p2.paragraph_format.space_after = Pt(4)

                    fg_left_sections = [
                        ("FG% - At the Rim",       "At Rim FG",       "At Rim Makes",       "At Rim Attempts"),
                        ("FG% - Paint 2s",         "In Paint FG",     "In Paint Makes",     "In Paint Attempts"),
                        ("FG% - Mid-Range 2s",     "Midrange 2s FG",  "Midrange 2s Makes",  "Midrange 2s Attempts"),
                        ("FG% - All 2pt Attempts", "Total 2pt FG",    "Total 2pt Makes",    "Total 2pt Attempts"),
                    ]

                    fg_right_sections = [
                        ("FG% – Above-the Break 3s", "Above Break 3s FG", "Above Break 3s Makes", "Above Break 3s Attempts"),
                        ("FG% – Corner 3s",          "Corner 3s FG",      "Corner 3s Makes",      "Corner 3s Attempts"),
                        ("FG% - All 3pt Attempts",   "Total 3pt FG",      "Total 3pt Makes",      "Total 3pt Attempts"),
                    ]

                    max_len_fg = max(len(fg_left_sections), len(fg_right_sections))

                    # ONE table for all FG% sections
                    fg_table = fg_doc.add_table(rows=max_len_fg, cols=2)
                    fg_table.autofit = True
                    remove_table_borders(fg_table)

                    for i in range(max_len_fg):
                        left_cell = fg_table.rows[i].cells[0]
                        right_cell = fg_table.rows[i].cells[1]

                        # ---------- LEFT SIDE (2PT) ----------
                        if i < len(fg_left_sections):
                            title, fg_col, make_col, att_col = fg_left_sections[i]
                            p = left_cell.paragraphs[0]
                            p.text = ""
                            r = p.add_run(title)
                            r.bold = True
                            r.font.size = Pt(16)
                            if header_color is not None:
                                r.font.color.rgb = header_color

                            if fg_col in df_fg.columns:
                                df_sorted = df_fg.sort_values(by=fg_col, ascending=False)
                                for rank, (_, row_fg) in enumerate(df_sorted.iterrows(), start=1):
                                    raw_val = row_fg.get(fg_col, 0.0)
                                    makes = int(row_fg.get(make_col, 0) or 0)
                                    attempts = int(row_fg.get(att_col, 0) or 0)

                                    if pd.isna(raw_val) or attempts == 0:
                                        val_display = 0.0
                                    else:
                                        val_display = float(raw_val)

                                    jersey = str(row_fg["Jersey"]).strip()
                                    name = str(row_fg["Player"])
                                    line = f"{rank}. #{jersey} {name} – {val_display:.1f}% ({makes}/{attempts})"
                                    left_cell.add_paragraph(line)

                        # ---------- RIGHT SIDE (3PT) ----------
                        if i < len(fg_right_sections):
                            title, fg_col, make_col, att_col = fg_right_sections[i]
                            p = right_cell.paragraphs[0]
                            p.text = ""
                            r = p.add_run(title)
                            r.bold = True
                            r.font.size = Pt(16)
                            if header_color is not None:
                                r.font.color.rgb = header_color

                            if fg_col in df_fg.columns:
                                df_sorted = df_fg.sort_values(by=fg_col, ascending=False)
                                for rank, (_, row_fg) in enumerate(df_sorted.iterrows(), start=1):
                                    raw_val = row_fg.get(fg_col, 0.0)
                                    makes = int(row_fg.get(make_col, 0) or 0)
                                    attempts = int(row_fg.get(att_col, 0) or 0)

                                    if pd.isna(raw_val) or attempts == 0:
                                        val_display = 0.0
                                    else:
                                        val_display = float(raw_val)

                                    jersey = str(row_fg["Jersey"]).strip()
                                    name = str(row_fg["Player"])
                                    line = f"{rank}. #{jersey} {name} – {val_display:.1f}% ({makes}/{attempts})"
                                    right_cell.add_paragraph(line)

                    # Force zero spacing for every paragraph inside the FG% table
                    for row in fg_table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(0)

                    fg_buffer = BytesIO()
                    fg_doc.save(fg_buffer)
                    fg_buffer.seek(0)

                    st.download_button(
                        "Download FG% DOCX",
                        data=fg_buffer,
                        file_name=f"{safe_team_name}_fg_percentages.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_fg_percentages",
                    )

else:
    st.info(
        "⬆️ Upload the CBB Analytics team player-profiles PDF to extract zone FGA% and generate Shot Diet + FG% DOCX."
    )





st.markdown(
    "<p style='text-align: center; font-size: 22px; color: gray;'>©Ari Jacobs 2025</p>",
    unsafe_allow_html=True,
)
